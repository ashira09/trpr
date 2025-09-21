from random import randint

from openpyxl import Workbook
from openpyxl.chart import BarChart, Reference 

import docx

# Генераций данных

data = [[f'Товар{i+1}' for i in range(10)], [randint(1,5) for i in range(10)], [round(randint(10000, 100000)/100, 2) for i in range(10)]]

# Операции с данными в Excel 

wb = Workbook()
sht = wb.active

sht.cell(row=1, column=1).value = 'Товар'
sht.cell(row=1, column=2).value = 'Количество'
sht.cell(row=1, column=3).value = 'Цена'
sht.cell(row=1, column=4).value = 'Стоимость'

for i in range(1,len(data[0]) + 1):
    for j in range(1,len(data) + 1):
        sht.cell(row=i, column=j).value = data[j-1][i-1]
    sht.cell(row=i, column=j+1).value = data[2][i-1] * data[1][i-1]

values = Reference(sht, min_col=3, min_row=2, max_col=3, max_row=10)
labels = Reference(sht, min_col=1, min_row=2, max_col=1, max_row=10)
chart = BarChart()
chart.add_data(values)
chart.set_categories(labels)
chart.title = "Стоимости товаров"
chart.x_axis.title = "Товары"
chart.y_axis.title = "Стоимость"
sht.add_chart(chart, "F1")

wb.save('data.xlsx')

# Операции с данными в Word

doc = docx.Document()

heading = doc.add_heading('Стоимости товаров', 0)

table = doc.add_table(cols=4, rows=11)

table.cell(row_idx=0, col_idx=0).add_paragraph(text='Товар')
table.cell(row_idx=0, col_idx=1).add_paragraph(text='Количество')
table.cell(row_idx=0, col_idx=2).add_paragraph(text='Цена')
table.cell(row_idx=0, col_idx=3).add_paragraph(text='Стоимость')

for i in range(1, len(data[0])):
    for j in range(0, len(data)):
        table.cell(row_idx=i, col_idx=j).add_paragraph(text=str(data[j][i]))
    table.cell(row_idx=i, col_idx=j+1).add_paragraph(text=str(data[2][i] * data[1][i]))

for i in range(1,len(data[0]) + 1):
    for j in range(1,len(data) + 1):
        table.cell(row_idx=i, col_idx=j).add_paragraph(text=str(sht.cell(row=i, column=j).value))
    table.cell(row_idx=i, col_idx=j+1).add_paragraph(text=str(sht.cell(row=i, column=j+1).value))

doc.save('report.docx')